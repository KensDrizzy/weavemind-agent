"""Lightweight local parser for text-like knowledge files.

This parser is the dependency-free MVP fallback. It keeps the Docling/Milvus
architecture open by emitting structured elements instead of raw character
chunks.
"""

from __future__ import annotations

import html
import mimetypes
import os
import re
from pathlib import Path

from bs4 import BeautifulSoup

from knowledge_rag.models import ParsedElement
from knowledge_rag.parsers.base import DocumentParser


class PlainTextDocumentParser(DocumentParser):
    provider = "plain-text"
    version = "plain-text@1"

    _TEXT_EXTENSIONS = {
        ".txt", ".md", ".markdown", ".html", ".htm", ".csv", ".tsv",
        ".json", ".yaml", ".yml", ".xml", ".rst", ".log",
    }

    def supports(self, file_path: str) -> bool:
        suffix = Path(file_path).suffix.lower()
        return suffix in self._TEXT_EXTENSIONS or Path(file_path).is_file()

    def parse(self, file_path: str) -> list[ParsedElement]:
        suffix = Path(file_path).suffix.lower()
        text = self._read_text(file_path)
        if suffix in {".html", ".htm"}:
            text = self._html_to_text(text)
        return self._text_to_elements(text)

    def _read_text(self, file_path: str) -> str:
        suffix = Path(file_path).suffix.lower()
        if suffix == ".pdf":
            extracted = self._try_extract_pdf(file_path)
            if extracted:
                return extracted
        if suffix == ".docx":
            extracted = self._try_extract_docx(file_path)
            if extracted:
                return extracted
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()

    @staticmethod
    def _try_extract_pdf(file_path: str) -> str:
        try:
            from pypdf import PdfReader
        except Exception:
            return ""
        try:
            reader = PdfReader(file_path)
            pages = []
            for index, page in enumerate(reader.pages, 1):
                text = page.extract_text() or ""
                if text.strip():
                    pages.append(f"\n\n[page {index}]\n{text}")
            return "\n".join(pages)
        except Exception:
            return ""

    @staticmethod
    def _try_extract_docx(file_path: str) -> str:
        try:
            from docx import Document
        except Exception:
            return ""
        try:
            doc = Document(file_path)
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        parts.append(" | ".join(cells))
            return "\n\n".join(parts)
        except Exception:
            return ""

    @staticmethod
    def _html_to_text(text: str) -> str:
        soup = BeautifulSoup(text, "html.parser")
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()
        return html.unescape(soup.get_text("\n"))

    @staticmethod
    def _text_to_elements(text: str) -> list[ParsedElement]:
        elements: list[ParsedElement] = []
        section_path: list[str] = []
        page_number = None
        buffer: list[str] = []

        def flush_paragraph():
            nonlocal buffer
            paragraph = "\n".join(buffer).strip()
            buffer = []
            if paragraph:
                elements.append(ParsedElement(
                    text=paragraph,
                    element_type="paragraph",
                    page_number=page_number,
                    section_path=list(section_path),
                ))

        for raw_line in text.splitlines():
            line = raw_line.strip()
            page_match = re.match(r"^\[page\s+(\d+)\]$", line, re.I)
            if page_match:
                flush_paragraph()
                page_number = int(page_match.group(1))
                continue
            if not line:
                flush_paragraph()
                continue

            heading = _heading_text(line)
            if heading:
                flush_paragraph()
                level, title = heading
                section_path = section_path[: max(level - 1, 0)] + [title]
                elements.append(ParsedElement(
                    text=title,
                    element_type="title",
                    page_number=page_number,
                    section_path=list(section_path),
                ))
                continue

            if _looks_like_table_row(line):
                flush_paragraph()
                elements.append(ParsedElement(
                    text=line,
                    element_type="table",
                    page_number=page_number,
                    section_path=list(section_path),
                ))
                continue

            if re.match(r"^(\s*[-*+]\s+|\s*\d+[.)]\s+)", raw_line):
                flush_paragraph()
                elements.append(ParsedElement(
                    text=line,
                    element_type="list",
                    page_number=page_number,
                    section_path=list(section_path),
                ))
                continue

            buffer.append(line)

        flush_paragraph()
        if not elements and text.strip():
            elements.append(ParsedElement(text=text.strip()))
        return elements


def _heading_text(line: str) -> tuple[int, str] | None:
    markdown = re.match(r"^(#{1,6})\s+(.+)$", line)
    if markdown:
        return len(markdown.group(1)), markdown.group(2).strip()
    numbered = re.match(r"^(\d+(?:\.\d+){0,5})[、.\s]+(.{2,120})$", line)
    if numbered:
        return numbered.group(1).count(".") + 1, numbered.group(2).strip()
    if len(line) <= 80 and re.match(r"^[一二三四五六七八九十]+[、.]\s*", line):
        return 1, re.sub(r"^[一二三四五六七八九十]+[、.]\s*", "", line).strip()
    return None


def _looks_like_table_row(line: str) -> bool:
    if "|" in line and line.count("|") >= 2:
        return True
    return "\t" in line and len(line.split("\t")) >= 3
